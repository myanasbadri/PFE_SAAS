"""
Insert 8 diagrams into the rapport with introductory text:
Chapter 3:
  1. Diag activité - Pipeline extraction     -> after Figure 3.6
  2. Diag activité - Conformité fiscale       -> after Tableau 3.3
  3. Diag séquence - Auth JWT                 -> in section 3.8
Chapter 4:
  4. Diag séquence - Inscription 5 étapes     -> after Figure 4.3
  5. Diag séquence - Extraction intelligente  -> after Figure 4.4
  6. Diag séquence - Conseiller IA            -> after Figure 4.10
  7. Diag séquence - Notifications            -> after Figure 4.17
  8. Diag états - Cycle de vie facture        -> after Figure 4.6
"""
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = r"C:\Users\MSI\Downloads\Qwen2.5-3B-Instruct (INT4)\invoice rapport FINAL ff (1) - avec listes.docx"
OUTPUT = r"C:\Users\MSI\Downloads\Qwen2.5-3B-Instruct (INT4)\invoice rapport FINAL v2.docx"

doc = Document(INPUT)

# ── Helpers ──
def insert_paragraph_after(paragraph):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)

def create_caption(paragraph, text):
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    run.italic = True
    pPr = paragraph._p.get_or_add_pPr()
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '200')
    spacing.set(qn('w:before'), '120')
    pPr.append(spacing)

def create_body_text(paragraph, text):
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    pPr = paragraph._p.get_or_add_pPr()
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')
    pPr.append(jc)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '120')
    spacing.set(qn('w:before'), '120')
    spacing.set(qn('w:line'), '360')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def add_image_para(paragraph, image_path, width_cm):
    run = paragraph.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    pPr = paragraph._p.get_or_add_pPr()
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '60')
    spacing.set(qn('w:before'), '120')
    pPr.append(spacing)

def insert_block(anchor, image_path, caption, intro, width_cm):
    """Insert intro + image + caption after anchor. Returns last paragraph."""
    blank = insert_paragraph_after(anchor)
    intro_p = insert_paragraph_after(blank)
    create_body_text(intro_p, intro)
    img_p = insert_paragraph_after(intro_p)
    add_image_para(img_p, image_path, width_cm)
    cap_p = insert_paragraph_after(img_p)
    create_caption(cap_p, caption)
    print(f"  Inserted: {caption[:70]}")
    return cap_p

# ── Find anchor paragraphs (body only, skip liste des figures area) ──
def find_body_caption(prefix, min_para=220):
    """Find a figure/tableau caption in the body (not in TOC/liste area)."""
    for i, p in enumerate(doc.paragraphs):
        if i < min_para:
            continue
        text = p.text.strip()
        if p.style.name.startswith("toc"):
            continue
        if text.startswith(prefix):
            return i
    return None

# ── Define all 8 insertions with their anchors ──
# We'll find them, then insert in REVERSE document order to avoid index shifts.

BASE = r"C:\Users\MSI\Desktop\PFE"

insertions = []

# 1. Pipeline extraction activity diagram → after Figure 3.6
idx = find_body_caption("Figure 3.6")
if idx:
    insertions.append((idx, {
        "image": f"{BASE}\\Diagramme_Activite_Pipeline_Extraction.png",
        "caption": "Figure 3.8 — Diagramme d'activité du pipeline d'extraction intelligente",
        "intro": (
            "Le diagramme d'activité suivant offre une vue complémentaire du pipeline d'extraction "
            "en modélisant le flux de contrôle complet, depuis le dépôt du fichier par l'utilisateur "
            "jusqu'à la persistance de la facture extraite. Il met en évidence les quatre couloirs "
            "de responsabilité (utilisateur, backend Flask, moteur IA, conformité) et illustre "
            "le mécanisme de branchement selon le type de fichier : les fichiers XML court-circuitent "
            "la chaîne LLM grâce à une extraction déterministe, tandis que les fichiers PDF et images "
            "traversent la chaîne de fallback complète (Claude → Ollama → regex). "
            "Les traitements parallèles — classification biens/services, calcul du score de confiance "
            "et détection de la devise — sont représentés par une barre de synchronisation (fork/join), "
            "reflétant leur exécution concurrente dans le code."
        ),
        "width": 15,
    }))
    print(f"Found Figure 3.6 at para {idx}")

# 2. Conformity activity diagram → after Tableau 3.3
idx = find_body_caption("Tableau 3.3")
if idx:
    insertions.append((idx, {
        "image": f"{BASE}\\Diagramme_Activite_Conformite.png",
        "caption": "Figure 3.9 — Diagramme d'activité de la vérification de conformité fiscale",
        "intro": (
            "Le diagramme d'activité ci-dessous détaille le flux d'exécution du moteur de conformité "
            "fiscale, qui applique séquentiellement les sept familles de contrôles présentées dans "
            "le tableau précédent. Chaque famille constitue une partition distincte dans le diagramme, "
            "reflétant la séparation des préoccupations dans l'implémentation. "
            "Le scoring repose sur un système de pénalités pondérées : une anomalie critique coûte "
            "dix points, une erreur cinq points et un avertissement deux points. "
            "Le score final, calculé sur une base de cent vingt points, détermine le statut de conformité : "
            "pleinement conforme, conforme avec avertissements ou non conforme. "
            "Ce mécanisme permet de distinguer les factures immédiatement soumissibles à la plateforme "
            "TTN de celles nécessitant une correction préalable."
        ),
        "width": 13,
    }))
    print(f"Found Tableau 3.3 at para {idx}")

# 3. Auth JWT sequence diagram → in section 3.8 (after last paragraph before Ch4)
# Find the heading "3.8 Architecture de sécurité" and insert after the last paragraph of that section
idx_38 = None
idx_ch4 = None
for i, p in enumerate(doc.paragraphs):
    if i < 220: continue
    text = p.text.strip()
    if "3.8 Architecture de sécurité" in text and p.style.name.startswith("Heading"):
        idx_38 = i
    if text.startswith("Chapitre 4") and p.style.name == "Heading 1":
        idx_ch4 = i
        break

if idx_38 and idx_ch4:
    # Insert before Chapter 4, i.e., after the last non-empty paragraph of section 3.8
    anchor_idx = idx_ch4 - 1
    # Find last non-empty para before ch4
    for j in range(idx_ch4 - 1, idx_38, -1):
        if doc.paragraphs[j].text.strip():
            anchor_idx = j
            break
    insertions.append((anchor_idx, {
        "image": f"{BASE}\\Diagramme_Sequence_Authentification_JWT.png",
        "caption": "Figure 3.10 — Diagramme de séquence de l'authentification JWT et de l'autorisation multi-tenant",
        "intro": (
            "Le diagramme de séquence suivant détaille le flux complet d'authentification et d'autorisation "
            "mis en œuvre par la plateforme. La première phase modélise le processus de connexion : "
            "l'utilisateur soumet ses identifiants, le serveur vérifie le mot de passe via bcrypt "
            "et génère un token JWT contenant l'identifiant utilisateur, le rôle global et une durée "
            "de validité de vingt-quatre heures. "
            "La seconde phase illustre le traitement d'une requête authentifiée traversant la chaîne "
            "de trois décorateurs : @token_required vérifie la signature et l'expiration du JWT, "
            "@org_member_required contrôle l'appartenance à l'organisation demandée (avec mise en cache Redis "
            "de cinq minutes pour optimiser les performances), et @role_required compare le rôle "
            "de l'utilisateur avec le niveau minimum requis selon la hiérarchie owner > admin > member. "
            "Cette architecture en couches garantit qu'aucune requête ne peut accéder aux données "
            "d'une organisation tierce, assurant ainsi l'isolation multi-tenant."
        ),
        "width": 15,
    }))
    print(f"Found section 3.8 anchor at para {anchor_idx}")

# 4. Registration sequence diagram → after Figure 4.3
idx = find_body_caption("Figure 4.3")
if idx:
    insertions.append((idx, {
        "image": f"{BASE}\\Diagramme_Sequence_Inscription.png",
        "caption": "Figure 4.4 — Diagramme de séquence de l'inscription en cinq étapes",
        "intro": (
            "Le diagramme de séquence ci-dessous modélise le flux complet de l'inscription, "
            "structurée en cinq étapes adaptées aux spécificités fiscales tunisiennes. "
            "Les quatre premières étapes collectent progressivement les informations côté client "
            "sans interaction serveur : le type d'activité (commercial, service, industriel, "
            "profession libérale ou artisanal), le type de personne (physique ou morale), "
            "les informations professionnelles optionnelles (raison sociale, matricule fiscal, "
            "téléphone, adresse) et les identifiants du compte. "
            "La cinquième étape présente un récapitulatif avant la soumission finale. "
            "Côté serveur, le traitement atomique enchaîne la vérification d'unicité de l'email, "
            "le hachage bcrypt du mot de passe, la création de l'utilisateur et de son organisation "
            "personnelle, l'enregistrement dans le journal d'audit et la génération du token JWT, "
            "permettant une connexion immédiate après l'inscription."
        ),
        "width": 14,
    }))
    print(f"Found Figure 4.3 at para {idx}")

# 5. Extraction sequence diagram → after Figure 4.4 (which is now the new inscription diagram)
# Actually, we should insert after the existing extraction interface figure
# Let's find it by looking for "Interface d'extraction intelligente" in captions
for i, p in enumerate(doc.paragraphs):
    if i < 220: continue
    text = p.text.strip()
    if "Interface d'extraction intelligente" in text and text.startswith("Figure"):
        insertions.append((i, {
            "image": f"{BASE}\\Diagramme_Sequence_Extraction.png",
            "caption": "Figure 4.5 — Diagramme de séquence de l'extraction intelligente de facture",
            "intro": (
                "Le diagramme de séquence suivant illustre le flux complet depuis le dépôt d'un fichier "
                "jusqu'à l'affichage des résultats d'extraction. Le traitement repose sur une architecture "
                "asynchrone : dès la réception du fichier, l'API crée une tâche Celery et retourne "
                "immédiatement un identifiant de tâche, permettant au frontend de suivre la progression "
                "par interrogation périodique (polling). "
                "Le worker Celery crée d'abord une facture provisoire avec le statut « en traitement », "
                "puis détermine la stratégie d'extraction selon le format du fichier. "
                "Les fichiers XML bénéficient d'une extraction déterministe directe, tandis que "
                "les fichiers PDF et images passent par l'extraction de texte (pdfplumber ou OCR Tesseract "
                "trilingue français-anglais-arabe) suivie de l'analyse sémantique par la chaîne de "
                "fallback LLM (Claude → Ollama → regex). "
                "Le post-traitement normalise les dates, corrige les quantités, fusionne les doublons "
                "et calcule les scores de confiance individuels par champ avant de soumettre les données "
                "au moteur de conformité fiscale."
            ),
            "width": 15,
        }))
        print(f"Found extraction interface figure at para {i}")
        break

# 6. AI Advisor sequence diagram → after Conseiller financier IA figure
for i, p in enumerate(doc.paragraphs):
    if i < 220: continue
    text = p.text.strip()
    if "Conseiller financier IA" in text and text.startswith("Figure"):
        insertions.append((i, {
            "image": f"{BASE}\\Diagramme_Sequence_Conseiller_IA.png",
            "caption": "Figure 4.11 — Diagramme de séquence du conseiller financier IA",
            "intro": (
                "Le diagramme de séquence ci-dessous détaille le fonctionnement du conseiller financier "
                "conversationnel. Lors de chaque interaction, le backend construit dynamiquement un contexte "
                "riche en agrégeant les données des deux cents dernières factures de l'organisation : "
                "montants totaux, répartition par fournisseur, tendances mensuelles sur six mois, "
                "statuts et devises utilisées. Ce contexte est injecté dans le prompt système du modèle "
                "de langage, accompagné de l'historique conversationnel (limité aux vingt derniers messages "
                "pour maîtriser la taille du contexte). "
                "Le mécanisme de fallback assure la disponibilité du service : si l'API Claude "
                "est indisponible (timeout, quota épuisé ou erreur), le système bascule automatiquement "
                "vers le modèle Ollama local. Un badge dans l'interface indique à l'utilisateur "
                "quel modèle a généré la réponse, garantissant la transparence du traitement."
            ),
            "width": 15,
        }))
        print(f"Found Conseiller IA figure at para {i}")
        break

# 7. State diagram (invoice lifecycle) → after Gestion des factures figure
for i, p in enumerate(doc.paragraphs):
    if i < 220: continue
    text = p.text.strip()
    if "Gestion des factures" in text and text.startswith("Figure") and "séquence" not in text.lower():
        insertions.append((i, {
            "image": f"{BASE}\\Diagramme_Etats_Cycle_Vie_Facture.png",
            "caption": "Figure 4.7 — Diagramme d'états du cycle de vie d'une facture",
            "intro": (
                "Le diagramme d'états suivant modélise le cycle de vie complet d'une facture dans "
                "le système, depuis son dépôt initial jusqu'à sa soumission à la plateforme nationale "
                "TTN El Fatoora. Après le téléversement, la facture entre dans un état composite "
                "« en traitement » regroupant les sous-états d'extraction de texte, d'analyse par IA, "
                "de classification et de contrôle de conformité. "
                "En sortie de traitement, deux chemins sont possibles selon le score de confiance : "
                "les factures à confiance élevée passent directement à l'état « vérifiée », tandis que "
                "celles dont le score est inférieur au seuil sont dirigées vers l'interface de vérification "
                "humaine (human-in-the-loop). "
                "Le diagramme illustre également le mécanisme de partage inter-organisationnel par QR code : "
                "une facture conforme peut être partagée via un code unique, puis importée par un autre "
                "utilisateur dans son organisation, créant une copie traçable avec vérification "
                "d'intégrité SHA-256."
            ),
            "width": 12,
        }))
        print(f"Found Gestion des factures figure at para {i}")
        break

# 8. Notifications sequence diagram → after Préférences de notification figure
for i, p in enumerate(doc.paragraphs):
    if i < 220: continue
    text = p.text.strip()
    if "Préférences de notification" in text and text.startswith("Figure"):
        insertions.append((i, {
            "image": f"{BASE}\\Diagramme_Sequence_Notifications.png",
            "caption": "Figure 4.19 — Diagramme de séquence du système de notifications",
            "intro": (
                "Le diagramme de séquence ci-dessous détaille l'architecture du système de notifications "
                "multi-canal. La première partie illustre le flux d'émission interne : lorsqu'un événement "
                "survient (extraction terminée, facture mise à jour, etc.), le service de notifications "
                "consulte les préférences de l'utilisateur pour déterminer les canaux à activer "
                "(in-app, email, SMS) et vérifier le respect des heures silencieuses. "
                "La notification in-app est persistée dans MongoDB avec un TTL de quatre-vingt-dix jours, "
                "tandis que l'email est acheminé via Gmail SMTP avec un template HTML personnalisé. "
                "La seconde partie modélise la consultation côté client : le frontend récupère les "
                "notifications paginées avec le compteur de non-lues, permettant à l'utilisateur "
                "de marquer individuellement ou collectivement les notifications comme lues. "
                "L'indexation composée (user_id, created_at) et (user_id, is_read) garantit "
                "des performances optimales même avec un volume élevé de notifications."
            ),
            "width": 15,
        }))
        print(f"Found Préférences de notification figure at para {i}")
        break

print(f"\nTotal insertions planned: {len(insertions)}")

# Sort by paragraph index in REVERSE order
insertions.sort(key=lambda x: x[0], reverse=True)

for para_idx, data in insertions:
    anchor = doc.paragraphs[para_idx]
    insert_block(anchor, data["image"], data["caption"], data["intro"], data["width"])

doc.save(OUTPUT)
print(f"\nSaved to: {OUTPUT}")
print("Done!")
