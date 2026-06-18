"""
Update rapport: replace tall diagrams with split versions (part1 + part2).
For each tall diagram, replace the existing image with part1 and insert part2 after.
"""

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
import os

DOCX = r"C:\Users\MSI\Downloads\Qwen2.5-3B-Instruct (INT4)\invoice rapport FINAL v2 vvvvvvvvvv.docx"
IMG_DIR = r"C:\Users\MSI\Desktop\PFE"

NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}

# Map: (caption_fragment, part1_img, part2_img, caption1_text, caption2_text)
SPLITS = [
    (
        "Figure 3.6",
        "Diagramme_Sequence_Extraction_Part1.png",
        "Diagramme_Sequence_Extraction_Part2.png",
        "Figure 3.6a \u2014 Séquence d'extraction (1/2) : dépôt, validation et extraction IA",
        "Figure 3.6b \u2014 Séquence d'extraction (2/2) : post-traitement, persistance et résultats",
    ),
    (
        "Figure 3.8",
        "Diagramme_Activite_Conformite_Part1.png",
        "Diagramme_Activite_Conformite_Part2.png",
        "Figure 3.8a \u2014 Activité de conformité fiscale (1/2) : familles 1 à 4",
        "Figure 3.8b \u2014 Activité de conformité fiscale (2/2) : familles 5 à 7 et scoring",
    ),
    (
        "Figure 3.10",
        "Diagramme_Sequence_Authentification_JWT_Part1.png",
        "Diagramme_Sequence_Authentification_JWT_Part2.png",
        "Figure 3.10a \u2014 Authentification JWT (1/2) : connexion",
        "Figure 3.10b \u2014 Authentification JWT (2/2) : autorisation multi-tenant",
    ),
    (
        "Figure 4.6",
        "Diagramme_Sequence_Extraction_Part1.png",
        "Diagramme_Sequence_Extraction_Part2.png",
        "Figure 4.6a \u2014 Séquence d'extraction intelligente (1/2) : dépôt et extraction IA",
        "Figure 4.6b \u2014 Séquence d'extraction intelligente (2/2) : post-traitement et résultats",
    ),
    (
        "Figure 4.16",
        "Diagramme_Sequence_Partage_QR_Part1.png",
        "Diagramme_Sequence_Partage_QR_Part2.png",
        "Figure 4.16a \u2014 Partage par QR code (1/2) : génération du code",
        "Figure 4.16b \u2014 Partage par QR code (2/2) : scan et importation",
    ),
    (
        "Figure 4.23",
        "Diagramme_Sequence_Notifications_Part1.png",
        "Diagramme_Sequence_Notifications_Part2.png",
        "Figure 4.23a \u2014 Notifications (1/2) : émission et consultation",
        "Figure 4.23b \u2014 Notifications (2/2) : marquage et préférences",
    ),
]

# Also split the new interface diagram
INTERFACE_SPLIT = (
    "Figure 3.x",
    "Diagramme_Interface_Client_Backend_Part1.png",
    "Diagramme_Interface_Client_Backend_Part2.png",
    "Figure 3.x(a) \u2014 Interface client-serveur (1/2) : couche client",
    "Figure 3.x(b) \u2014 Interface client-serveur (2/2) : couche serveur",
)


def _find_caption_para(doc, caption_fragment):
    """Find paragraph index containing the caption text (skip TOC entries)."""
    # Search from the end to skip TOC entries; also verify a drawing exists nearby
    candidates = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith(caption_fragment):
            candidates.append(i)
    # Return the LAST match (body caption, not TOC)
    return candidates[-1] if candidates else None


def _find_image_para_before_caption(doc, caption_idx):
    """Find the image paragraph just before the caption."""
    for j in range(caption_idx - 1, max(caption_idx - 4, 0), -1):
        para = doc.paragraphs[j]
        if para._element.findall(".//w:drawing", NSMAP):
            return j
    return None


def _replace_image_blob(doc, para, img_path):
    """Replace the image data in a paragraph."""
    blips = para._element.findall(".//a:blip", NSMAP)
    for blip in blips:
        r_embed = blip.get(f"{{{NSMAP['r']}}}embed")
        if r_embed:
            rel = doc.part.rels[r_embed]
            with open(img_path, "rb") as f:
                rel.target_part._blob = f.read()
            return True
    return False


def _make_caption_element(doc, text):
    """Create a centered, italic caption paragraph element."""
    body = doc.element.body
    c_para = etree.SubElement(body, f"{{{NSMAP['w']}}}p")
    c_ppr = etree.SubElement(c_para, f"{{{NSMAP['w']}}}pPr")
    c_jc = etree.SubElement(c_ppr, f"{{{NSMAP['w']}}}jc")
    c_jc.set(f"{{{NSMAP['w']}}}val", "center")
    c_run = etree.SubElement(c_para, f"{{{NSMAP['w']}}}r")
    c_rpr = etree.SubElement(c_run, f"{{{NSMAP['w']}}}rPr")
    etree.SubElement(c_rpr, f"{{{NSMAP['w']}}}i")
    c_sz = etree.SubElement(c_rpr, f"{{{NSMAP['w']}}}sz")
    c_sz.set(f"{{{NSMAP['w']}}}val", "22")
    c_t = etree.SubElement(c_run, f"{{{NSMAP['w']}}}t")
    c_t.text = text
    c_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    body.remove(c_para)
    return c_para


def _make_image_element(doc, img_path):
    """Create an image paragraph element."""
    body = doc.element.body
    temp_para = doc.add_paragraph()
    run = temp_para.runs[0] if temp_para.runs else temp_para.add_run()
    run.add_picture(img_path, width=Cm(16))
    temp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elem = temp_para._element
    body.remove(elem)
    return elem


def process_split(doc, caption_frag, part1_img, part2_img, caption1, caption2):
    """Replace one tall diagram with two parts."""
    part1_path = os.path.join(IMG_DIR, part1_img)
    part2_path = os.path.join(IMG_DIR, part2_img)

    if not os.path.exists(part1_path) or not os.path.exists(part2_path):
        print(f"  SKIP {caption_frag}: images not found")
        return False

    # Find the caption paragraph
    caption_idx = _find_caption_para(doc, caption_frag)
    if caption_idx is None:
        print(f"  SKIP {caption_frag}: caption not found")
        return False

    # Find the image paragraph before the caption
    img_idx = _find_image_para_before_caption(doc, caption_idx)
    if img_idx is None:
        print(f"  SKIP {caption_frag}: image para not found before caption")
        return False

    img_para = doc.paragraphs[img_idx]
    caption_para = doc.paragraphs[caption_idx]

    print(f"  {caption_frag}: img at para {img_idx}, caption at para {caption_idx}")

    # Step 1: Replace existing image with part1
    _replace_image_blob(doc, img_para, part1_path)

    # Step 2: Update existing caption text to caption1
    caption_para.clear()
    run = caption_para.add_run(caption1)
    run.italic = True
    run.font.size = Cm(0.39)  # ~11pt
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Step 3: Insert part2 image + caption AFTER the existing caption
    ref_element = caption_para._element
    body = doc.element.body

    # Create part2 image element
    img2_elem = _make_image_element(doc, part2_path)
    # Create part2 caption element
    cap2_elem = _make_caption_element(doc, caption2)

    # Insert after caption: image2 then caption2
    ref_element.addnext(cap2_elem)
    ref_element.addnext(img2_elem)

    print(f"  OK: Split into {caption1[:50]}... + {caption2[:50]}...")
    return True


def main():
    print(f"Loading: {DOCX}")
    doc = Document(DOCX)

    print("\n=== Splitting tall diagrams ===")
    for entry in SPLITS:
        process_split(doc, *entry)

    # Handle the interface diagram (Figure 3.x)
    print("\n=== Splitting interface diagram ===")
    process_split(doc, *INTERFACE_SPLIT)

    out_path = DOCX.replace(".docx", " SPLIT.docx")
    print(f"\nSaving: {out_path}")
    doc.save(out_path)
    print("Done!")


if __name__ == "__main__":
    main()
