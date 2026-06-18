"""
Add 3 new diagrams to the rapport:
1. Diagramme de sequence - Edition d'une facture     -> after Figure 4.6 in section 4.3.4
2. Diagramme de sequence - Suppression d'une facture -> after diagram 1
3. Diagramme de sequence - Partage QR code           -> after Figure 4.9 in section 4.4.1
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = r"C:\Users\MSI\Downloads\Qwen2.5-3B-Instruct (INT4)\invoice rapport FINAL ff (1) - avec listes.docx"
OUTPUT = r"C:\Users\MSI\Downloads\Qwen2.5-3B-Instruct (INT4)\invoice rapport FINAL ff (1) - avec listes.docx"

doc = Document(INPUT)

def insert_paragraph_after(paragraph):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)

def create_caption(paragraph, caption_text):
    run = paragraph.add_run(caption_text)
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    run.italic = True
    pPr = paragraph._p.get_or_add_pPr()
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '200')
    spacing.set(qn('w:before'), '120')
    pPr.append(spacing)

def create_intro_text(paragraph, text):
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    pPr = paragraph._p.get_or_add_pPr()
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')
    pPr.append(jc)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '120')
    spacing.set(qn('w:before'), '120')
    spacing.set(qn('w:line'), '276')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def add_image_paragraph(paragraph, image_path, width_cm):
    run = paragraph.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    pPr = paragraph._p.get_or_add_pPr()
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '60')
    spacing.set(qn('w:before'), '120')
    pPr.append(spacing)

def insert_diagram(anchor, image_path, caption_text, intro_text, width_cm):
    """Insert intro + image + caption after anchor. Returns the caption paragraph."""
    blank = insert_paragraph_after(anchor)
    intro_p = insert_paragraph_after(blank)
    create_intro_text(intro_p, intro_text)
    img_p = insert_paragraph_after(intro_p)
    add_image_paragraph(img_p, image_path, width_cm)
    cap_p = insert_paragraph_after(img_p)
    create_caption(cap_p, caption_text)
    print(f"  Inserted: {caption_text}")
    return cap_p

# ── Find insertion points ──
# We process in reverse document order to avoid index shifts

# 1. Find "Figure 4.9" caption (for sharing diagram - insert after it)
fig49_idx = None
# 2. Find "Figure 4.6" caption (for edit & delete diagrams - insert after it)
fig46_idx = None

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style.startswith("toc"):
        continue
    # Skip liste des figures entries (they have tab + page number, and are before para 220)
    if "\t" in text and i < 220:
        continue
    if text.startswith("Figure 4.9"):
        fig49_idx = i
        print(f"Found Figure 4.9 at para {i}: {text}")
    if text.startswith("Figure 4.6"):
        fig46_idx = i
        print(f"Found Figure 4.6 at para {i}: {text}")

# ── Insert in reverse order (last first) ──

# 3. Insert sharing diagram after Figure 4.9
if fig49_idx:
    anchor = doc.paragraphs[fig49_idx]
    insert_diagram(
        anchor,
        r"C:\Users\MSI\Desktop\PFE\Diagramme_Sequence_Partage_QR.png",
        "Figure 4.10 — Diagramme de séquence du partage et de l'importation par QR code",
        "Le diagramme de séquence ci-dessous détaille le flux complet de partage inter-organisationnel par QR code. "
        "Il couvre trois phases : la génération du code de partage unique (format INV-XXXXXXXX) et du QR code côté client, "
        "la consultation de la facture partagée par le destinataire via scan ou saisie manuelle, "
        "et enfin l'importation de la facture dans l'organisation cible avec les contrôles de sécurité associés "
        "(prévention de l'auto-importation et des doublons).",
        15
    )
    print()

# 1 & 2. Insert edit then delete diagrams after Figure 4.6
if fig46_idx:
    anchor = doc.paragraphs[fig46_idx]

    # Insert edit diagram first
    last_p = insert_diagram(
        anchor,
        r"C:\Users\MSI\Desktop\PFE\Diagramme_Sequence_Edition_Facture.png",
        "Figure 4.7 — Diagramme de séquence de l'édition d'une facture",
        "Le diagramme de séquence suivant illustre le processus complet de modification d'une facture. "
        "Après authentification JWT et vérification de l'appartenance organisationnelle, "
        "le système valide les champs modifiés (données, tags, statut), "
        "effectue la mise à jour atomique dans MongoDB, puis enregistre l'événement dans le journal d'audit "
        "et émet une notification adaptée au type de modification.",
        15
    )

    # Insert delete diagram after edit
    insert_diagram(
        last_p,
        r"C:\Users\MSI\Desktop\PFE\Diagramme_Sequence_Suppression_Facture.png",
        "Figure 4.8 — Diagramme de séquence de la suppression d'une facture",
        "Le diagramme ci-dessous présente le flux de suppression d'une facture. "
        "Après la confirmation par l'utilisateur via un dialogue dédié, "
        "le système vérifie l'authentification et le contrôle d'accès organisationnel, "
        "puis procède à une suppression définitive (hard delete) dans MongoDB. "
        "L'opération est tracée dans le journal d'audit avec le nom de fichier original pour référence ultérieure.",
        15
    )

doc.save(OUTPUT)
print(f"\nSaved to: {OUTPUT}")
print("Done! 3 new diagrams added to the rapport.")
