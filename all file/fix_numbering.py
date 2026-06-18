"""
Fix figure numbering after inserting new diagrams.
Renumber all Figure 4.x captions sequentially (both in liste des figures and document body).
"""
import re
from docx import Document

INPUT  = r"C:\Users\MSI\Downloads\Qwen2.5-3B-Instruct (INT4)\invoice rapport FINAL ff (1) - avec listes.docx"
OUTPUT = r"C:\Users\MSI\Downloads\Qwen2.5-3B-Instruct (INT4)\invoice rapport FINAL ff (1) - avec listes.docx"

doc = Document(INPUT)

# Step 1: Collect all Figure 4.x captions in document order (body only, skip liste des figures)
# The liste des figures is roughly paras 172-205, body starts after ~220
body_figures = []  # list of (para_index, old_number, caption_suffix)

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if i < 220:
        continue  # Skip liste des figures area
    if p.style.name.startswith("toc"):
        continue

    m = re.match(r'^(Figure 4\.\d+)\s*(.*)', text)
    if m:
        old_label = m.group(1)  # "Figure 4.7"
        suffix = m.group(2)     # "— Gestion des factures"
        old_num = old_label.split('.')[1]  # "7"
        body_figures.append((i, old_label, suffix))

print("=== Body figures in order ===")
for idx, (para_i, old_label, suffix) in enumerate(body_figures):
    new_num = idx + 1  # 4.1, 4.2, ...
    new_label = f"Figure 4.{new_num}"
    print(f"  Para {para_i}: {old_label} -> {new_label} {suffix[:60]}")

# Step 2: Build renaming map
# body_figures are already in order, so new number = position in list + 1
rename_map = {}  # (para_index) -> (old_text_prefix, new_text_prefix)
for idx, (para_i, old_label, suffix) in enumerate(body_figures):
    new_num = idx + 1
    new_label = f"Figure 4.{new_num}"
    if old_label != new_label:
        rename_map[para_i] = (old_label, new_label)

print(f"\n=== {len(rename_map)} figures need renumbering ===")

# Step 3: Apply renaming in document body
for i, p in enumerate(doc.paragraphs):
    if i in rename_map:
        old_label, new_label = rename_map[i]
        old_text = p.text
        # Update runs
        for run in p.runs:
            if old_label in run.text:
                run.text = run.text.replace(old_label, new_label)
                print(f"  Body para {i}: {old_label} -> {new_label}")
                break

# Step 4: Now rebuild the Liste des figures entries
# Find all entries in the liste (they have tab + page number, in paras ~173-205)
liste_figures = []
liste_start = None
liste_end = None

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text == "Liste des figures":
        liste_start = i
    if text == "Liste des tableaux":
        liste_end = i
        break
    if liste_start and i > liste_start:
        m = re.match(r'^(Figure 4\.\d+)', text)
        if m:
            liste_figures.append((i, m.group(1), text))

print(f"\n=== Liste des figures: Figure 4.x entries (paras {liste_start}-{liste_end}) ===")

# Build a map from body figure order to their new numbers
# We need to match liste entries to body entries by their suffix
body_new_labels = {}
for idx, (para_i, old_label, suffix) in enumerate(body_figures):
    new_num = idx + 1
    new_label = f"Figure 4.{new_num}"
    # Extract the descriptive part after the dash
    desc = re.sub(r'^[—–-]\s*', '', suffix).split('\t')[0].strip()
    body_new_labels[desc] = new_label

# Update liste des figures entries
for para_i, old_label, full_text in liste_figures:
    p = doc.paragraphs[para_i]
    # Extract description from the entry
    # Format: "Figure 4.X — Description\tPageNum"
    desc_match = re.match(r'^Figure 4\.\d+\s*[—–-]\s*(.*?)(\t\d+)?$', full_text)
    if desc_match:
        desc = desc_match.group(1).strip()
        # Find matching new label
        matched = False
        for body_desc, new_label in body_new_labels.items():
            if desc in body_desc or body_desc in desc:
                if old_label != new_label:
                    for run in p.runs:
                        if old_label in run.text:
                            run.text = run.text.replace(old_label, new_label)
                            print(f"  Liste para {para_i}: {old_label} -> {new_label} ({desc[:50]})")
                            matched = True
                            break
                else:
                    matched = True
                if matched:
                    break
        if not matched:
            print(f"  WARNING: No match for liste entry: {full_text[:80]}")

doc.save(OUTPUT)
print(f"\nSaved to: {OUTPUT}")
print("Done!")
