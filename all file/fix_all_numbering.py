"""
Fix ALL figure numbering after inserting 8 new diagrams.
Renumber sequentially within each chapter, update both body and Liste des figures.
Also update the Liste des figures to include all new entries.
"""
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = r"C:\Users\MSI\Downloads\Qwen2.5-3B-Instruct (INT4)\invoice rapport FINAL v2.docx"
OUTPUT = r"C:\Users\MSI\Downloads\Qwen2.5-3B-Instruct (INT4)\invoice rapport FINAL v2.docx"

doc = Document(INPUT)

# ── Step 1: Collect ALL figure captions in body (skip TOC & liste area) ──
body_figures = []  # (para_idx, chapter, old_label, suffix_after_number)

for i, p in enumerate(doc.paragraphs):
    if i < 230:
        continue
    text = p.text.strip()
    if p.style.name.startswith("toc"):
        continue
    m = re.match(r'^(Figure (\d+)\.(\d+))\s*(.*)', text)
    if m:
        old_label = m.group(1)  # "Figure 3.8"
        chapter = int(m.group(2))
        suffix = m.group(4)      # "— ..."
        body_figures.append((i, chapter, old_label, suffix))

# ── Step 2: Assign new sequential numbers per chapter ──
chapter_counters = {}
renaming = []  # (para_idx, old_label, new_label, full_new_caption)

for para_idx, chapter, old_label, suffix in body_figures:
    if chapter not in chapter_counters:
        chapter_counters[chapter] = 0
    chapter_counters[chapter] += 1
    new_num = chapter_counters[chapter]
    new_label = f"Figure {chapter}.{new_num}"
    full_caption = f"{new_label} {suffix}"
    renaming.append((para_idx, old_label, new_label, full_caption))

print("=== New figure numbering ===")
for para_idx, old_label, new_label, full_caption in renaming:
    changed = " *" if old_label != new_label else ""
    print(f"  Para {para_idx}: {old_label} -> {new_label}{changed}  {full_caption[len(new_label):len(new_label)+50]}")

# ── Step 3: Apply renaming in body ──
changes_made = 0
for para_idx, old_label, new_label, _ in renaming:
    if old_label == new_label:
        continue
    p = doc.paragraphs[para_idx]
    for run in p.runs:
        if old_label in run.text:
            run.text = run.text.replace(old_label, new_label)
            changes_made += 1
            break

print(f"\n{changes_made} body captions renamed")

# ── Step 4: Rebuild the entire Liste des figures ──
# Find the liste boundaries
liste_start = None
liste_end = None
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text == "Liste des figures":
        liste_start = i
    if text == "Liste des tableaux" and liste_start:
        liste_end = i
        break

if liste_start and liste_end:
    # Delete all existing entries between "Liste des figures" and "Liste des tableaux"
    # We'll clear their text and then add new ones
    entries_to_clear = []
    for i in range(liste_start + 1, liste_end):
        text = doc.paragraphs[i].text.strip()
        if text.startswith("Figure"):
            entries_to_clear.append(i)

    # Clear old entries
    for idx in entries_to_clear:
        p = doc.paragraphs[idx]
        for run in p.runs:
            run.text = ""
        # Remove tab stops etc
        p._p.clear()

    print(f"\nCleared {len(entries_to_clear)} old liste entries")

    # Build page number mapping from TOC
    toc_map = {}
    for p in doc.paragraphs:
        if p.style.name.startswith("toc"):
            text = p.text.strip()
            if "\t" in text:
                parts = text.rsplit("\t", 1)
                try:
                    page = int(parts[1].strip())
                    toc_map[parts[0].strip()] = page
                except ValueError:
                    pass

    # Build heading→page index for the body
    heading_pages = []
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading"):
            heading_text = p.text.strip()
            for toc_key, page in toc_map.items():
                if heading_text in toc_key or toc_key in heading_text:
                    heading_pages.append((i, page))
                    break
            else:
                m2 = re.match(r'(\d+\.\d+(?:\.\d+)?)', heading_text)
                if m2:
                    section = m2.group(1)
                    for toc_key, page in toc_map.items():
                        if section in toc_key:
                            heading_pages.append((i, page))
                            break

    def get_page(para_idx):
        best = 1
        for h_idx, h_page in heading_pages:
            if h_idx <= para_idx:
                best = h_page
            else:
                break
        return best

    # Now insert new entries after "Liste des figures" heading
    def insert_paragraph_after(paragraph):
        new_p = OxmlElement('w:p')
        paragraph._p.addnext(new_p)
        from docx.text.paragraph import Paragraph
        return Paragraph(new_p, paragraph._parent)

    def create_list_entry(paragraph, text, page_num):
        pPr = paragraph._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:leader'), 'dot')
        tab.set(qn('w:pos'), '9072')
        tabs.append(tab)
        pPr.append(tabs)
        run = paragraph.add_run(text)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        paragraph.add_run("\t")
        run3 = paragraph.add_run(str(page_num))
        run3.font.size = Pt(12)
        run3.font.name = "Times New Roman"
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '60')
        spacing.set(qn('w:before'), '60')
        spacing.set(qn('w:line'), '276')
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)

    # Insert all figures in REVERSE order (since each insert goes right after anchor)
    anchor = doc.paragraphs[liste_start]
    all_entries = []
    for para_idx, _, new_label, full_caption in renaming:
        page = get_page(para_idx)
        caption_text = f"{new_label} {full_caption[len(new_label):]}"
        # Clean up: remove tab and page number if present in caption
        caption_text = caption_text.split("\t")[0].strip()
        all_entries.append((caption_text, page))

    # Insert in reverse so they end up in correct order
    for caption_text, page in reversed(all_entries):
        new_p = insert_paragraph_after(anchor)
        create_list_entry(new_p, caption_text, page)

    print(f"Inserted {len(all_entries)} new liste entries")

doc.save(OUTPUT)
print(f"\nSaved to: {OUTPUT}")
print("Done!")
