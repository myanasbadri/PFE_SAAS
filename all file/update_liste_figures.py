"""
Update the Liste des figures to include the 3 new diagrams.
The new figures are inserted in section 4.3.4 and 4.4.1.
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = r"C:\Users\MSI\Downloads\Qwen2.5-3B-Instruct (INT4)\invoice rapport FINAL ff (1) - avec listes.docx"
OUTPUT = r"C:\Users\MSI\Downloads\Qwen2.5-3B-Instruct (INT4)\invoice rapport FINAL ff (1) - avec listes.docx"

doc = Document(INPUT)

def insert_paragraph_after(paragraph):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)

def create_list_entry(paragraph, text, page_num):
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '9072')
    tabs.append(tab)
    pPr.append(tabs)
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    paragraph.add_run("\t")
    run3 = paragraph.add_run(str(page_num))
    run3.font.size = Pt(12)
    run3.font.name = "Times New Roman"
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '60')
    spacing.set(qn('w:before'), '60')
    spacing.set(qn('w:line'), '276')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

# Find the "Figure 4.6" entry in the Liste des figures
# and add the 3 new entries after it
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    # We're looking in the Liste des figures area (before para 220)
    if i > 220:
        break
    if "Figure 4.6" in text and "\t" in text and i < 220:
        print(f"Found Figure 4.6 entry at para {i}: {text}")
        anchor = p

        # Insert Figure 4.7 and 4.8 after Figure 4.6 entry
        entry1 = insert_paragraph_after(anchor)
        create_list_entry(entry1,
            "Figure 4.7 — Diagramme de séquence de l'édition d'une facture", 38)

        entry2 = insert_paragraph_after(entry1)
        create_list_entry(entry2,
            "Figure 4.8 — Diagramme de séquence de la suppression d'une facture", 39)

        print("  Added Figure 4.7 and 4.8 entries")
        break

# Now find the "Figure 4.9" entry and add Figure 4.10 after it
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if i > 230:
        break
    if "Figure 4.9" in text and "\t" in text and i < 230:
        print(f"Found Figure 4.9 entry at para {i}: {text}")
        anchor = p

        entry = insert_paragraph_after(anchor)
        create_list_entry(entry,
            "Figure 4.10 — Diagramme de séquence du partage et de l'importation par QR code", 42)

        print("  Added Figure 4.10 entry")
        break

doc.save(OUTPUT)
print(f"\nSaved to: {OUTPUT}")
print("Done!")
