"""
Add 'Liste des figures' and 'Liste des tableaux' to the Word document,
inserted after the Table of Contents and before the Introduction.
"""
import re
import copy
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = r"C:\Users\MSI\Downloads\Qwen2.5-3B-Instruct (INT4)\invoice rapport FINAL ff (1).docx"
OUTPUT = r"C:\Users\MSI\Downloads\Qwen2.5-3B-Instruct (INT4)\invoice rapport FINAL ff (1) - avec listes.docx"

doc = Document(INPUT)

# ── 1. Parse the TOC to build a heading→page mapping ──────────────────────
toc_map = {}  # heading_text → page_number
for p in doc.paragraphs:
    if p.style.name.startswith("toc"):
        text = p.text.strip()
        # TOC entries look like "Heading text\tPageNumber"
        if "\t" in text:
            parts = text.rsplit("\t", 1)
            heading_text = parts[0].strip()
            try:
                page = int(parts[1].strip())
                toc_map[heading_text] = page
            except ValueError:
                pass

print(f"Parsed {len(toc_map)} TOC entries")

# ── 2. Build heading index for the document ──────────────────────────────
heading_pages = []  # list of (paragraph_index, page_number)
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith("Heading"):
        heading_text = p.text.strip()
        # Try to match with TOC
        for toc_key, page in toc_map.items():
            if heading_text in toc_key or toc_key in heading_text:
                heading_pages.append((i, page))
                break
        else:
            # Try partial match on section numbers
            m = re.match(r'(\d+\.\d+(?:\.\d+)?)', heading_text)
            if m:
                section = m.group(1)
                for toc_key, page in toc_map.items():
                    if section in toc_key:
                        heading_pages.append((i, page))
                        break

def get_page_for_para(para_idx):
    """Find the page number for a paragraph by looking at the nearest preceding heading."""
    best_page = 1
    for h_idx, h_page in heading_pages:
        if h_idx <= para_idx:
            best_page = h_page
        else:
            break
    return best_page

# ── 3. Collect figure and table captions ─────────────────────────────────
figures = []
tables = []

figure_pattern = re.compile(r'^Figure\s+\d+\.\d+\s*[—–-]')
table_pattern  = re.compile(r'^Tableau\s+\d+\.\d+\s*[—–-]')

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    # Skip TOC entries
    if p.style.name.startswith("toc"):
        continue
    # Skip headings that just mention "Tableau" (like "4.3.3 Tableau de bord")
    if p.style.name.startswith("Heading"):
        continue

    if figure_pattern.match(text):
        page = get_page_for_para(i)
        figures.append((text, page))
    elif table_pattern.match(text):
        page = get_page_for_para(i)
        tables.append((text, page))

print(f"Found {len(figures)} figures and {len(tables)} tables")

# ── 4. Find insertion point (after last TOC entry, before Introduction) ──
# Find the last TOC entry paragraph
last_toc_idx = 0
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith("toc"):
        last_toc_idx = i

# Find first empty paragraph after TOC (we'll use that area to insert)
insert_after_idx = last_toc_idx
for i in range(last_toc_idx + 1, len(doc.paragraphs)):
    if doc.paragraphs[i].text.strip() == "":
        insert_after_idx = i
        break

print(f"Inserting after paragraph {insert_after_idx}")

# ── 5. Helper: create a formatted entry (caption + dotted tab + page) ────
def create_list_entry(paragraph, text, page_num):
    """Format a list entry like a TOC entry with dotted tab leader."""
    # Clear the paragraph
    for run in paragraph.runs:
        run.text = ""

    # Set paragraph style similar to TOC
    pPr = paragraph._p.get_or_add_pPr()

    # Add tab stop at right margin with dot leader
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '9072')  # ~16cm from left margin
    tabs.append(tab)
    pPr.append(tabs)

    # Add the text
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # Add tab character
    run2 = paragraph.add_run("\t")

    # Add page number
    run3 = paragraph.add_run(str(page_num))
    run3.font.size = Pt(12)
    run3.font.name = "Times New Roman"

    # Set spacing
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '60')
    spacing.set(qn('w:before'), '60')
    spacing.set(qn('w:line'), '276')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def create_heading(paragraph, text):
    """Format a heading for the list sections."""
    for run in paragraph.runs:
        run.text = ""

    run = paragraph.add_run(text)
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    run.bold = True

    pPr = paragraph._p.get_or_add_pPr()

    # Center alignment
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)

    # Spacing
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '240')
    spacing.set(qn('w:before'), '360')
    spacing.set(qn('w:line'), '276')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def insert_paragraph_after(paragraph):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)

# ── 6. Insert the lists ──────────────────────────────────────────────────
# We need to insert in reverse order since each insert goes right after the anchor

# Start from the insertion point
anchor = doc.paragraphs[insert_after_idx]

# We'll build all paragraphs to insert, then insert them in order
# Insert: blank, "Liste des figures" heading, figure entries...,
#          blank, "Liste des tableaux" heading, table entries..., blank

# Insert a page break before the lists
new_p = insert_paragraph_after(anchor)
run = new_p.add_run()
br = OxmlElement('w:br')
br.set(qn('w:type'), 'page')
run._r.append(br)
anchor = new_p

# ── Liste des figures ──
heading_p = insert_paragraph_after(anchor)
create_heading(heading_p, "Liste des figures")
anchor = heading_p

for text, page in figures:
    entry_p = insert_paragraph_after(anchor)
    create_list_entry(entry_p, text, page)
    anchor = entry_p

# Page break between the two lists
new_p = insert_paragraph_after(anchor)
run = new_p.add_run()
br = OxmlElement('w:br')
br.set(qn('w:type'), 'page')
run._r.append(br)
anchor = new_p

# ── Liste des tableaux ──
heading_p = insert_paragraph_after(anchor)
create_heading(heading_p, "Liste des tableaux")
anchor = heading_p

for text, page in tables:
    entry_p = insert_paragraph_after(anchor)
    create_list_entry(entry_p, text, page)
    anchor = entry_p

# Add a page break after the lists (before Introduction)
new_p = insert_paragraph_after(anchor)
run = new_p.add_run()
br = OxmlElement('w:br')
br.set(qn('w:type'), 'page')
run._r.append(br)

# ── 7. Save ──────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"\nSaved to: {OUTPUT}")
print("Done!")
